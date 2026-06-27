from __future__ import annotations

import io
import logging
import os
import tempfile
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

import aiofiles
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.concurrency import run_in_threadpool
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse

from .ai_service import analyze_case, stream_chat

logger = logging.getLogger(__name__)
from .demo_data import build_demo_case, get_all_cases, get_case_summaries
from .models import (
    AnalyzeJobCreated,
    AnalyzeJobStatus,
    AnalyzeRequest,
    CaseAnalysis,
    CaseSummary,
    ChatRequest,
    FetchUrlRequest,
)
from .ocr_adapter import MistralOcrAdapter, PptxAdapter, PypdfAdapter, XlsxAdapter
from .stripe_service import create_maxx_checkout_session, stripe_configured
from .ocr_models import OcrInput

app = FastAPI(title="SchedaPRO API", version="1.0.0")

_DEFAULT_ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:5178",
    "http://127.0.0.1:5178",
    "http://localhost:5179",
    "http://127.0.0.1:5179",
    "http://localhost:5180",
    "http://127.0.0.1:5180",
    "http://localhost:5181",
    "http://127.0.0.1:5181",
    "https://localhost",
    "http://localhost",
    "capacitor://localhost",
    "https://digitaltrainer.netlify.app",
    "https://schedapro-838.netlify.app",
]
_EXTRA_ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv("ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]
_ALLOWED_ORIGINS = [*_DEFAULT_ALLOWED_ORIGINS, *_EXTRA_ALLOWED_ORIGINS]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "Accept"],
)


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok", "service": "schedapro-api", "version": "1.0.0"}


# ── Cases list ───────────────────────────────────────────────────────────────

@app.get("/api/cases", response_model=list[CaseSummary])
def list_cases(lang: str = "it") -> list[CaseSummary]:
    return get_case_summaries(lang)


@app.get("/api/cases/{case_id}", response_model=CaseAnalysis)
def get_case(case_id: str, lang: str = "it") -> CaseAnalysis:
    cases = get_all_cases(lang)
    if case_id not in cases:
        raise HTTPException(status_code=404, detail=f"Case '{case_id}' not found")
    return cases[case_id]


# ── Legacy demo endpoint (kept for backward compat) ──────────────────────────

@app.get("/api/demo-case", response_model=CaseAnalysis)
def get_demo_case() -> CaseAnalysis:
    return build_demo_case()


# ── AI analysis ──────────────────────────────────────────────────────────────

@app.post("/api/analyze-text", response_model=CaseAnalysis)
def analyze_text(request: AnalyzeRequest) -> CaseAnalysis:
    """Run AI analysis on provided text materials."""
    logger.info("analyze-text: title=%s, materials=%d, mode=%s, lang=%s",
                request.case_title, len(request.materials), request.mode, request.language)
    try:
        return analyze_case(request)
    except ValueError as exc:
        # Model-level issues (truncation, invalid JSON) — surface the real message
        msg = str(exc)
        logger.error("analyze-text value error: %s", msg)
        raise HTTPException(status_code=422, detail=msg) from exc
    except Exception as exc:
        logger.error("analyze-text failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Analisi non disponibile. Riprova tra qualche secondo. "
                   "Se il problema persiste, prova con meno documenti o in modalità Pro."
        ) from exc


# ── Background analysis jobs ─────────────────────────────────────────────────
# The analysis can take a while; the client must be able to navigate away, lock
# the phone, or refresh and still pick the result back up. So we run it as a
# fire-and-poll job: POST starts it and returns a job_id immediately; the client
# polls GET until status is done/error. Jobs live in-process (fine for a single
# Render instance; a cold start mid-job loses it — the client then re-submits).

_ANALYSIS_GENERIC_ERROR = (
    "Analisi non disponibile. Riprova tra qualche secondo. "
    "Se il problema persiste, prova con meno documenti o in modalità Pro."
)
_JOB_TTL_SECONDS = 60 * 60  # keep finished jobs around an hour for late pollers
_jobs: dict[str, dict[str, Any]] = {}
_jobs_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="analyze-job")


def _purge_stale_jobs() -> None:
    cutoff = time.time() - _JOB_TTL_SECONDS
    for jid in [j for j, v in _jobs.items() if v.get("updated_at", 0) < cutoff]:
        _jobs.pop(jid, None)


def _run_analysis_job(job_id: str, request: AnalyzeRequest) -> None:
    try:
        result = analyze_case(request)
        _jobs[job_id] = {"status": "done", "result": result, "error": None, "updated_at": time.time()}
    except ValueError as exc:
        logger.error("analyze-job %s value error: %s", job_id, exc)
        _jobs[job_id] = {"status": "error", "result": None, "error": str(exc), "updated_at": time.time()}
    except Exception as exc:
        logger.error("analyze-job %s failed: %s", job_id, exc, exc_info=True)
        _jobs[job_id] = {"status": "error", "result": None, "error": _ANALYSIS_GENERIC_ERROR, "updated_at": time.time()}


@app.post("/api/analyze-jobs", response_model=AnalyzeJobCreated)
def create_analysis_job(request: AnalyzeRequest) -> AnalyzeJobCreated:
    """Accept an analysis, run it in the background, return a job id to poll."""
    _purge_stale_jobs()
    job_id = uuid.uuid4().hex
    _jobs[job_id] = {"status": "running", "result": None, "error": None, "updated_at": time.time()}
    logger.info("analyze-job %s queued: title=%s, materials=%d, mode=%s",
                job_id, request.case_title, len(request.materials), request.mode)
    _jobs_executor.submit(_run_analysis_job, job_id, request)
    return AnalyzeJobCreated(job_id=job_id)


@app.get("/api/analyze-jobs/{job_id}", response_model=AnalyzeJobStatus)
def get_analysis_job(job_id: str) -> AnalyzeJobStatus:
    """Poll an analysis job. 404 if unknown (e.g. lost to a cold start)."""
    job = _jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job di analisi sconosciuto o scaduto.")
    return AnalyzeJobStatus(status=job["status"], result=job["result"], error=job["error"])


# ── Chat (SSE streaming) ─────────────────────────────────────────────────────

@app.post("/api/chat")
def chat_endpoint(request: ChatRequest) -> StreamingResponse:
    """Stream a chat response from Claude as Server-Sent Events."""
    try:
        return StreamingResponse(
            stream_chat(request),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )
    except Exception as exc:
        logger.error("chat failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail="Chat non disponibile. Riprova tra qualche secondo.") from exc


# ── File upload ───────────────────────────────────────────────────────────────

_pypdf = PypdfAdapter()
_mistral = MistralOcrAdapter()
_pptx = PptxAdapter()
_xlsx = XlsxAdapter()

MAX_UPLOAD_BYTES = int(os.environ.get("PLT_MAX_UPLOAD_BYTES", str(50 * 1024 * 1024)))
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PPTX_MIME = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _archive_upload_response(filename: str, mime: str, archive_type: str, size_bytes: int) -> dict[str, Any]:
    if archive_type == "zip":
        extracted_text = "[File ZIP rilevato. Estrai i file e caricali singolarmente: l'estrazione automatica non è supportata per sicurezza.]"
        engine = "archive-zip"
        warnings = ["I file ZIP non vengono aperti automaticamente. Estrai e carica i singoli file."]
    else:
        extracted_text = "[File RAR rilevato. Estrai i file e caricali singolarmente: il formato RAR non è supportato.]"
        engine = "archive-rar"
        warnings = ["Formato RAR non supportato. Estrai e carica i singoli file (PDF, DOCX, immagini, ecc.)."]

    return {
        "upload_id": str(uuid.uuid4()),
        "filename": filename,
        "mime_type": mime,
        "size_bytes": size_bytes,
        "extracted_text": extracted_text,
        "engine": engine,
        "warnings": warnings,
        "status": "needs_ocr",
    }


def _result_to_upload_payload(
    *,
    filename: str,
    mime: str,
    size_bytes: int,
    extracted_text: str,
    engine: str,
    warnings: list[str],
    ready: bool,
) -> dict[str, Any]:
    return {
        "upload_id": str(uuid.uuid4()),
        "filename": filename,
        "mime_type": mime,
        "size_bytes": size_bytes,
        "extracted_text": extracted_text,
        "engine": engine,
        "warnings": warnings,
        "status": "ready" if ready else "needs_ocr",
    }


def _extract_docx_from_path(file_path: Path) -> tuple[str, str, list[str], bool]:
    try:
        from docx import Document  # type: ignore

        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n\n".join(paragraphs)
        return text or "[Documento Word vuoto]", "python-docx", [], bool(text)
    except Exception as exc:
        return f"[Errore estrazione DOCX: {exc}]", "python-docx-error", [str(exc)], False


def _extract_text_from_path(file_path: Path) -> tuple[str, str, list[str], bool]:
    text = file_path.read_text(encoding="utf-8", errors="replace")
    return text, "passthrough", [], bool(text.strip())


def _ocr_result_to_text(result, success_label: str, failure_prefix: str) -> tuple[str, str, list[str], bool]:
    warnings = [w.message for w in result.warnings]
    if result.success:
        extracted_text = "\n\n".join(f"[{success_label} {p.page}]\n{p.text}" for p in result.pages)
        return extracted_text, result.engine, warnings, True
    extracted_text = f"[{failure_prefix}: {warnings[-1] if warnings else 'Incolla il testo manualmente.'}]"
    return extracted_text, result.engine, warnings, False


def _extract_pptx_from_path(file_path: Path, mime: str) -> tuple[str, str, list[str], bool]:
    result = _pptx.extract(OcrInput(file_path=file_path, mime_type=mime))
    return _ocr_result_to_text(result, "Slide", "Estrazione PPTX non riuscita")


def _extract_xlsx_from_path(file_path: Path, mime: str) -> tuple[str, str, list[str], bool]:
    result = _xlsx.extract(OcrInput(file_path=file_path, mime_type=mime))
    warnings = [w.message for w in result.warnings]
    if result.success:
        extracted_text = "\n\n".join(
            f"[{p.text.split(chr(10))[0]}]\n" + "\n".join(p.text.split(chr(10))[1:]) for p in result.pages
        )
        return extracted_text, result.engine, warnings, True
    return f"[Estrazione XLSX non riuscita: {warnings[-1] if warnings else 'Prova a convertire in PDF e ricaricare.'}]", result.engine, warnings, False


def _extract_pdf_or_ocr_from_path(file_path: Path, mime: str) -> tuple[str, str, list[str], bool]:
    ocr_input = OcrInput(file_path=file_path, mime_type=mime)
    result = _pypdf.extract(ocr_input)
    if not result.success:
        result = _mistral.extract(ocr_input)
    return _ocr_result_to_text(result, "Pagina", f"Estrazione non riuscita per {mime}")


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)) -> dict[str, Any]:
    """Extract uploaded file text without blocking the event loop on parsing work."""
    mime = file.content_type or ""
    filename = file.filename or "documento"
    lower_filename = filename.lower()

    if file.size is not None and file.size > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File troppo grande per l'upload alpha")

    is_text = mime.startswith("text/") or lower_filename.endswith((".txt", ".rtf", ".csv"))
    is_docx = mime == _DOCX_MIME or lower_filename.endswith(".docx")
    is_pptx = mime == _PPTX_MIME or lower_filename.endswith(".pptx")
    is_xlsx = mime == _XLSX_MIME or lower_filename.endswith(".xlsx")
    is_zip = lower_filename.endswith(".zip")
    is_rar = lower_filename.endswith(".rar")

    if is_zip:
        return _archive_upload_response(filename, mime, "zip", file.size or 0)
    if is_rar:
        return _archive_upload_response(filename, mime, "rar", file.size or 0)

    fd, temp_name = tempfile.mkstemp(prefix="plt-upload-")
    os.close(fd)
    temp_path = Path(temp_name)
    size_bytes = 0

    try:
        async with aiofiles.open(temp_path, "wb") as out:
            while chunk := await file.read(1024 * 1024):
                size_bytes += len(chunk)
                if size_bytes > MAX_UPLOAD_BYTES:
                    raise HTTPException(status_code=413, detail="File troppo grande per l'upload alpha")
                await out.write(chunk)

        if is_text:
            extracted_text, engine, warnings, ready = await run_in_threadpool(_extract_text_from_path, temp_path)
        elif is_docx:
            extracted_text, engine, warnings, ready = await run_in_threadpool(_extract_docx_from_path, temp_path)
        elif is_pptx:
            extracted_text, engine, warnings, ready = await run_in_threadpool(_extract_pptx_from_path, temp_path, mime)
        elif is_xlsx:
            extracted_text, engine, warnings, ready = await run_in_threadpool(_extract_xlsx_from_path, temp_path, mime)
        else:
            extracted_text, engine, warnings, ready = await run_in_threadpool(_extract_pdf_or_ocr_from_path, temp_path, mime)

        return _result_to_upload_payload(
            filename=filename,
            mime=mime,
            size_bytes=size_bytes,
            extracted_text=extracted_text,
            engine=engine,
            warnings=warnings,
            ready=ready,
        )
    finally:
        try:
            temp_path.unlink(missing_ok=True)
        except OSError:
            logger.warning("Could not delete temporary upload file: %s", temp_path)


# ── URL fetch ─────────────────────────────────────────────────────────────────

@app.post("/api/fetch-url")
async def fetch_url_content(request: FetchUrlRequest) -> dict[str, Any]:
    """Fetch a URL and extract its text content for use as a case material."""
    import re as _re
    url = request.url.strip()
    if not _re.match(r"^https?://", url):
        raise HTTPException(status_code=422, detail="URL non valido: sono supportati solo http:// e https://")

    label = request.name.strip() or url

    try:
        resp = await run_in_threadpool(
            _fetch_url_text, url, label
        )
        return resp
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("fetch-url failed for %s: %s", url, exc, exc_info=True)
        raise HTTPException(status_code=502, detail=f"Impossibile recuperare il contenuto dall'URL: {exc}") from exc


def _fetch_url_text(url: str, label: str) -> dict[str, Any]:
    import httpx

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        )
    }
    with httpx.Client(follow_redirects=True, timeout=15) as client:
        r = client.get(url, headers=headers)
        r.raise_for_status()
    html = r.text

    extracted_text: str | None = None
    engine = "trafilatura"
    try:
        import trafilatura  # type: ignore
        extracted_text = trafilatura.extract(html, include_comments=False, include_tables=True)
    except Exception:
        pass

    if not extracted_text:
        engine = "beautifulsoup"
        try:
            from bs4 import BeautifulSoup  # type: ignore
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            extracted_text = soup.get_text(separator="\n", strip=True)
        except Exception:
            pass

    if not extracted_text:
        extracted_text = ""

    status = "ready" if extracted_text.strip() else "empty"
    warnings: list[str] = []
    if status == "empty":
        warnings.append("Nessun testo estratto dall'URL. Il contenuto potrebbe essere dinamico (JavaScript) o non accessibile.")

    import re as _re
    from urllib.parse import urlparse
    parsed = urlparse(url)
    filename = label or (parsed.netloc + parsed.path).rstrip("/").replace("/", "_") or "documento-web"
    filename = _re.sub(r"[^\w\-.]", "_", filename)[:80]

    return {
        "upload_id": str(uuid.uuid4()),
        "filename": filename,
        "mime_type": "text/plain",
        "size_bytes": len(extracted_text.encode("utf-8")),
        "extracted_text": extracted_text,
        "engine": engine,
        "warnings": warnings,
        "status": status,
        "source_url": url,
    }


# ── Voice transcription ───────────────────────────────────────────────────────

@app.post("/api/transcribe")
async def transcribe_audio(file: UploadFile = File(...)) -> dict[str, Any]:
    """Transcribe audio via Groq Whisper. Accepts webm, mp4, mp3, wav, ogg, m4a."""
    import os
    groq_key = os.environ.get("GROQ_API_KEY")
    if not groq_key:
        raise HTTPException(status_code=503, detail="GROQ_API_KEY non configurata")

    content = await file.read()
    filename = file.filename or "audio.webm"

    from groq import Groq  # type: ignore
    client = Groq(api_key=groq_key)

    transcription = client.audio.transcriptions.create(
        file=(filename, content),
        model="whisper-large-v3-turbo",
        language="it",
        response_format="text",
    )

    return {"text": transcription if isinstance(transcription, str) else transcription.text}


# ── Brief → DOCX export ───────────────────────────────────────────────────────

from pydantic import BaseModel as _BaseModel

class BriefExportRequest(_BaseModel):
    case_title: str
    brief_markdown: str

@app.post("/api/export-brief")
def export_brief(req: BriefExportRequest) -> Response:
    """Convert brief_markdown to a DOCX file and return it for download."""
    import re as _re
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_par = doc.add_heading(req.case_title, level=0)
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _brief_contains_any(text: str, needles: tuple[str, ...]) -> bool:
        lowered = text.lower()
        return any(n in lowered for n in needles)

    is_english = _brief_contains_any(req.brief_markdown, ("case summary", "defense attorney"))
    disclaimer = (
        "WORKING DOCUMENT — generated by Pocket Legal Triage with DeepSeek V4. "
        "For DEFENSE ATTORNEY review and editing only."
        if is_english
        else "DOCUMENTO DI LAVORO — generato da Pocket Legal Triage con DeepSeek V4. "
             "Riservato alla verifica e modifica del DIFENSORE."
    )
    p = doc.add_paragraph()
    run = p.add_run(disclaimer)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    for line in req.brief_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue
        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        # List items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, stripped[2:])
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = _re.sub(r'[^\w\-]', '_', req.case_title)[:60]
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )

def _add_inline(paragraph, text: str) -> None:
    """Add a paragraph run with basic bold/italic support."""
    import re as _re
    parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ── Maxx subscription checkout (Stripe) ────────────────────────────────────────

class CheckoutRequest(_BaseModel):
    email: str | None = None


@app.post("/api/checkout")
def create_checkout(req: CheckoutRequest) -> dict[str, str]:
    """Create a Stripe Checkout Session for the Maxx plan and return its URL.

    Returns 503 when Stripe isn't configured (no keys) so the frontend can fall
    back to its placeholder note instead of breaking.
    """
    if not stripe_configured():
        raise HTTPException(status_code=503, detail="Checkout non ancora disponibile.")
    try:
        url = create_maxx_checkout_session(req.email)
    except Exception as exc:  # noqa: BLE001 — surface a clean 502 to the client
        logger.error("create_checkout: Stripe session failed: %s", exc)
        raise HTTPException(status_code=502, detail="Errore nella creazione del checkout.") from exc
    return {"url": url}
