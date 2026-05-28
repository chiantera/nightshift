import io
from fastapi.testclient import TestClient
from docx import Document
from app.main import app

def test_export_brief_italian():
    client = TestClient(app)
    req_data = {
        "case_title": "Fascicolo Rossi",
        "brief_markdown": "### Sintesi del caso\nQuesto è un promemoria per l'udienza del difensore."
    }
    response = client.post("/api/export-brief", json=req_data)
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # Read the docx
    doc = Document(io.BytesIO(response.content))
    assert len(doc.paragraphs) > 0
    
    # Check disclaimer in paragraphs
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "DOCUMENTO DI LAVORO" in full_text
    assert "DIFENSORE" in full_text
    assert "DeepSeek V4" in full_text
    assert "Pocket Legal Triage" in full_text

def test_export_brief_english():
    client = TestClient(app)
    req_data = {
        "case_title": "Rossi Case",
        "brief_markdown": "### Case Summary\nThis is a memo for the defense attorney."
    }
    response = client.post("/api/export-brief", json=req_data)
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # Read the docx
    doc = Document(io.BytesIO(response.content))
    assert len(doc.paragraphs) > 0
    
    # Check disclaimer in paragraphs
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "WORKING DOCUMENT" in full_text
    assert "DEFENSE ATTORNEY" in full_text
    assert "DeepSeek V4" in full_text
    assert "Pocket Legal Triage" in full_text
